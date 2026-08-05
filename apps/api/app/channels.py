"""Self-service SMS channel activation: DLT registration (either self-declared with a
certificate upload, auto-approved instantly, or a request for Textzi to handle registration
on the customer's behalf, which goes into an admin review queue since real DLT registration is
a manual submission to the telecom registry) and channel subscription payment. Mirrors the
Razorpay dev/real dual-path pattern already used in wallet.py/payments.py."""
import hmac
import io
import os
import re
import secrets
import uuid
from datetime import datetime, timedelta, timezone

import razorpay
from docx import Document
from fastapi import APIRouter, Depends, File, Form, HTTPException, Response, UploadFile
from sqlalchemy import select
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session

from .auth import _send_platform_otp_sms, require_recent_2fa, require_user
from .config import settings
from .database import get_db
from .email_service import render_email, send_email
from .invoicing import create_draft_invoice, issue_invoice
from .permissions import require_capability
from .models import (
    ApiKey, ApiKeyActionOtp, ChannelSettings, ChannelSubscription, DltOnboardingRequest, DltOnboardingRequestDocument, Entity, Header,
    Message, PaymentOrder, PeId, Status, Template, User,
)
from .schemas import (
    ApiKeyCreateRequest, ApiKeyCreateResponse, ApiKeyIpWhitelistUpdate, ApiKeyOtpRequest, ApiKeyOtpResponse, ApiKeyOut,
    ApiKeyRevokeRequest, ChannelSettingsOut, ChannelSettingsUpdate, ChannelStatusResponse, DltDocumentOut,
    DltOnboardingRequestOut, DltRequestQuoteResponse, RazorpayOrderResponse, RazorpayVerifyRequest, RechargeResponse,
    ReportsSummaryResponse,
)
from .security import decrypt_secret, encrypt_secret, generate_otp, hash_api_key, hash_otp
from .services import GST_RATE, DomainError, channel_active, mask_aadhar, mask_email, mask_mobile, require_channel_active, resolve_channel_fees, resolve_user_entity

API_KEY_OTP_TTL_MINUTES = 10
API_KEY_OTP_MAX_ATTEMPTS = 5
ACTION_LABELS = {
    "generate_key": "generating a new API key", "update_ip_whitelist": "updating an API key's IP allow-list",
    "revoke_key": "revoking an API key",
}


def _issue_api_key_otp(db: Session, user: User, action: str) -> ApiKeyOtpResponse:
    code = generate_otp()
    expires_at = datetime.now(timezone.utc) + timedelta(minutes=API_KEY_OTP_TTL_MINUTES)
    db.add(ApiKeyActionOtp(user_id=user.id, action=action, code_hash=hash_otp(code), expires_at=expires_at))
    db.commit()
    # Mobile first if verified -- but only report success if the provider actually accepted the
    # send (_send_platform_otp_sms's own return value), not just that we attempted it. Platform
    # SMS being unconfigured or the platform wallet being empty must never leave the customer
    # staring at a code that will never arrive with no way to know it failed -- fall back to the
    # channel that always works instead.
    dev_code = code if settings.environment == "development" else None
    if user.mobile and user.mobile_verified and _send_platform_otp_sms(db, user.mobile, code):
        return ApiKeyOtpResponse(sent_via="mobile", masked_destination=mask_mobile(user.mobile), dev_otp_code=dev_code)
    send_email(
        db, to=user.email, subject="Your Textzi verification code",
        html_body=render_email(
            "Verify this action",
            f"<p>Hi {user.full_name},</p><p>Use the code below to confirm {ACTION_LABELS.get(action, 'this action')} on your Textzi account.</p>"
            f"<p style=\"margin:24px 0; text-align:center;\"><span style=\"display:inline-block; font-size:28px; font-weight:bold; letter-spacing:6px; color:#1a1a1a; background:#f4f4f5; padding:12px 24px; border-radius:6px;\">{code}</span></p>"
            f"<p>This code expires in {API_KEY_OTP_TTL_MINUTES} minutes. If you didn't request this, someone may have access to your account -- change your password and contact support.</p>",
        ),
    )
    return ApiKeyOtpResponse(sent_via="email", masked_destination=mask_email(user.email), dev_otp_code=dev_code)


def _consume_api_key_otp(db: Session, user_id: str, action: str, code: str) -> None:
    record = db.scalar(
        select(ApiKeyActionOtp)
        .where(ApiKeyActionOtp.user_id == user_id, ApiKeyActionOtp.action == action, ApiKeyActionOtp.consumed_at.is_(None))
        .order_by(ApiKeyActionOtp.created_at.desc()),
    )
    if not record:
        raise HTTPException(status_code=422, detail="No pending verification code for this action; request a new one")
    if record.expires_at < datetime.now(timezone.utc):
        raise HTTPException(status_code=422, detail="Verification code has expired; request a new one")
    if record.attempts >= API_KEY_OTP_MAX_ATTEMPTS:
        raise HTTPException(status_code=429, detail="Too many attempts; request a new code")
    record.attempts += 1
    if not hmac.compare_digest(record.code_hash, hash_otp(code)):
        db.commit()
        raise HTTPException(status_code=422, detail="Incorrect verification code")
    record.consumed_at = datetime.now(timezone.utc)
    db.commit()

router = APIRouter(prefix="/v1/channels/sms", tags=["channels"])

MAX_UPLOAD_BYTES = 10 * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".jpg", ".jpeg", ".png"}


def _save_upload(upload: UploadFile, subdir: str) -> tuple[str, bytes]:
    ext = os.path.splitext(upload.filename or "")[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=422, detail=f"Unsupported file type '{ext}'. Allowed: PDF, JPG, PNG.")
    content = upload.file.read()
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=422, detail="File is too large (max 10MB).")
    directory = os.path.join(settings.uploads_dir, subdir)
    os.makedirs(directory, exist_ok=True)
    stored_name = f"{uuid.uuid4()}{ext}"
    stored_path = os.path.join(directory, stored_name)
    with open(stored_path, "wb") as f:
        f.write(content)
    return stored_path, content


def _dlt_request_out(r: DltOnboardingRequest, docs: list[DltOnboardingRequestDocument]) -> DltOnboardingRequestOut:
    return DltOnboardingRequestOut(
        id=r.id, status=r.status, notes=r.notes,
        company_name=r.company_name, company_pan=r.company_pan, company_gst=r.company_gst,
        authorized_signatory_name=r.authorized_signatory_name, contact_number=r.contact_number, contact_email=r.contact_email,
        authorized_person_aadhar_masked=mask_aadhar(decrypt_secret(r.authorized_person_aadhar_encrypted)) if r.authorized_person_aadhar_encrypted else None,
        total_amount=float(r.total_amount), created_at=r.created_at.isoformat(),
        documents=[DltDocumentOut(id=d.id, filename=d.filename, document_type=d.document_type) for d in docs],
    )


def _build_authorization_letter_sample() -> bytes:
    """A fillable template matching Tata's own required declaration-letter wording -- the
    customer prints it on their company letterhead, fills in the blanks by hand, signs it, and
    uploads the scanned/photographed result back via the authorization_letter upload field."""
    doc = Document()
    note = doc.add_paragraph()
    note.add_run("[Print this letter on your Company Letterhead]").italic = True
    doc.add_paragraph()
    doc.add_paragraph("Date: _______________")
    doc.add_paragraph()
    doc.add_paragraph("To,")
    doc.add_paragraph("Tata Teleservices Ltd.,")
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.add_run("Sub: Declaration of the Authorized Person for Telemarketing Activities").bold = True
    doc.add_paragraph()
    doc.add_paragraph(
        "I, _____________________ (Name of the authorized signatory), Director / Partner of "
        "_____________________ (Company Name) (should be Director as per MCA / assigned by Board "
        "of Directors in MOA, or legal partner as per valid Partnership deed, etc.) have the "
        "authority to sign on behalf of our company for Telemarketer/Enterprise related activities "
        "OR am authorizing _____________________ (Authorized person's name) to sign on behalf of "
        "the company for related activities.",
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "We further confirm that the entity is liable for and bound by all acts of commission and "
        "omission by the authorized representative. All acts committed by the above authorized "
        "representative shall be treated as if these acts were committed by the entity.",
    )
    doc.add_paragraph()
    doc.add_paragraph("Authorized Person's Name: _____________________")
    doc.add_paragraph("Contact Number: _____________________")
    doc.add_paragraph("Email ID: _____________________")
    doc.add_paragraph("Authorized ID Proof Type: _____________________")
    doc.add_paragraph("ID Proof Number: _____________________")
    doc.add_paragraph()
    doc.add_paragraph("Thanks & Regards,")
    doc.add_paragraph()
    doc.add_paragraph("Authorized signatory's / person's name & Employee ID: _____________________")
    doc.add_paragraph("Signature of Partner/Director/Authorized Person: _____________________")
    doc.add_paragraph("Name of Signatory: _____________________")
    doc.add_paragraph("Company Stamp / Seal: _____________________")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _razorpay_client() -> razorpay.Client:
    if not settings.razorpay_key_id or not settings.razorpay_key_secret:
        raise HTTPException(status_code=503, detail="Razorpay is not configured. Set RAZORPAY_KEY_ID and RAZORPAY_KEY_SECRET in .env.")
    return razorpay.Client(auth=(settings.razorpay_key_id, settings.razorpay_key_secret))


def _dlt_status(db: Session, entity_id: str) -> str:
    has_active_pe = db.scalar(select(PeId).where(PeId.entity_id == entity_id, PeId.status == Status.active).limit(1)) is not None
    if has_active_pe:
        return "approved"
    pending_request = db.scalar(
        select(DltOnboardingRequest)
        .where(DltOnboardingRequest.entity_id == entity_id, DltOnboardingRequest.status.in_(["paid", "in_progress"]))
        .limit(1),
    )
    if pending_request:
        return "pending_review"
    return "not_started"


@router.get("/status", response_model=ChannelStatusResponse)
def get_channel_status(user: User = Depends(require_user), db: Session = Depends(get_db)):
    try:
        entity = resolve_user_entity(db, user)
        fees = resolve_channel_fees(db, "sms")
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    subscription = db.get(ChannelSubscription, (entity.id, "sms"))
    subscription_price = float(fees.subscription_price)
    subscription_paid = subscription_price == 0 or (subscription is not None and subscription.paid_at is not None)
    return ChannelStatusResponse(
        subscription_price=subscription_price,
        subscription_paid=subscription_paid,
        dlt_status=_dlt_status(db, entity.id),
        channel_active=channel_active(db, entity.id, "sms"),
    )


@router.get("/settings", response_model=ChannelSettingsOut)
def get_channel_settings(user: User = Depends(require_recent_2fa), db: Session = Depends(get_db)):
    try:
        entity = resolve_user_entity(db, user)
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    row = db.get(ChannelSettings, (entity.id, "sms"))
    return ChannelSettingsOut(encryption_enabled=bool(row and row.encryption_enabled), dr_webhook_url=row.dr_webhook_url if row else None)


@router.put("/settings", response_model=ChannelSettingsOut)
def update_channel_settings(payload: ChannelSettingsUpdate, user: User = Depends(require_recent_2fa), db: Session = Depends(get_db)):
    """`dr_webhook_url` is where Textzi relays a message's delivery report once one arrives from
    the upstream provider -- Tata itself never sees or calls this URL directly; Textzi is always
    the one calling out to it (see webhooks.py)."""
    try:
        entity = resolve_user_entity(db, user)
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    row = db.get(ChannelSettings, (entity.id, "sms"))
    if not row:
        row = ChannelSettings(entity_id=entity.id, channel="sms")
        db.add(row)
    row.encryption_enabled = payload.encryption_enabled
    row.dr_webhook_url = payload.dr_webhook_url
    db.commit()
    return ChannelSettingsOut(encryption_enabled=row.encryption_enabled, dr_webhook_url=row.dr_webhook_url)


@router.post("/subscription/recharge", response_model=RechargeResponse)
def simulate_subscription_payment(user: User = Depends(require_user), db: Session = Depends(get_db)):
    try:
        entity = resolve_user_entity(db, user)
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    subscription = db.get(ChannelSubscription, (entity.id, "sms"))
    if not subscription:
        subscription = ChannelSubscription(entity_id=entity.id, channel="sms")
        db.add(subscription)
    subscription.paid_at = datetime.now(timezone.utc)
    fees = resolve_channel_fees(db, "sms")
    if float(fees.subscription_price) > 0:
        gst_amount = round(float(fees.subscription_price) * GST_RATE, 2)
        invoice = create_draft_invoice(db, entity, type="channel_subscription", base_amount=float(fees.subscription_price), gst_amount=gst_amount)
        issue_invoice(db, invoice)
    db.commit()
    return RechargeResponse(entity_id=entity.id, amount=0, available_balance=0, dev_note="Development mode — channel subscription marked paid directly.")


@router.get("/dlt/quote", response_model=DltRequestQuoteResponse)
def quote_dlt_request(user: User = Depends(require_user), db: Session = Depends(get_db)):
    try:
        resolve_user_entity(db, user)
        fees = resolve_channel_fees(db, "sms")
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    combined = float(fees.dlt_platform_fee) + float(fees.dlt_service_fee)
    gst_amount = round(combined * GST_RATE, 2)
    return DltRequestQuoteResponse(combined_fee=combined, gst_amount=gst_amount, total_amount=round(combined + gst_amount, 2))


@router.post("/dlt/self-service")
def submit_self_service_dlt(
    pe_value: str = Form(...),
    operator: str = Form(...),
    header_id: str = Form(...),
    header_value: str = Form(...),
    certificate: UploadFile = File(...),
    user: User = Depends(require_user),
    db: Session = Depends(get_db),
):
    """A customer who already has their own DLT registration self-declares it here, uploads
    the PE certificate as proof, and it's approved immediately -- no admin review, matching how
    admin-created DLT assets already go active on creation without an approval step."""
    try:
        entity = resolve_user_entity(db, user)
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    if db.scalar(select(PeId).where(PeId.entity_id == entity.id, PeId.value == pe_value)):
        raise HTTPException(status_code=409, detail="This PE ID is already registered for your account")
    stored_path, _ = _save_upload(certificate, f"pe-certificates/{entity.id}")
    pe = PeId(entity_id=entity.id, value=pe_value, operator=operator, certificate_path=stored_path, status=Status.active)
    db.add(pe); db.flush()
    header = Header(pe_id=pe.id, header_id=header_id, value=header_value, status=Status.active)
    db.add(header)
    db.commit(); db.refresh(pe); db.refresh(header)
    return {"pe_id": pe.id, "header_id": header.id, "status": "approved"}


PAN_PATTERN = re.compile(r"[A-Z]{5}[0-9]{4}[A-Z]")
AADHAR_PATTERN = re.compile(r"[0-9]{12}")
MOBILE_PATTERN = re.compile(r"[6-9][0-9]{9}")


@router.get("/dlt/authorization-letter-sample")
def download_authorization_letter_sample(user: User = Depends(require_user)):
    return Response(
        content=_build_authorization_letter_sample(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=Textzi-Authorization-Letter-Sample.docx"},
    )


@router.post("/dlt/request", response_model=DltOnboardingRequestOut)
def submit_dlt_request(
    company_name: str = Form(...),
    company_pan: str = Form(...),
    authorized_signatory_name: str = Form(...),
    contact_number: str = Form(...),
    contact_email: str = Form(...),
    authorized_person_aadhar: str = Form(...),
    company_gst: str | None = Form(default=None),
    notes: str | None = Form(default=None),
    documents: list[UploadFile] = File(...),
    authorization_letter: UploadFile = File(...),
    user: User = Depends(require_user),
    db: Session = Depends(get_db),
):
    """A customer without existing DLT registration asks Textzi to handle it for them. Fees are
    snapshotted now so later admin fee-config changes never retroactively alter a submitted
    request. This only files the request -- payment happens via a separate endpoint. The KYC
    fields mirror exactly what Tata's own telemarketer-registration process requires (company
    PAN/GST, authorized signatory, and a signed authorization letter declaring who's allowed to
    act on the company's behalf) -- collecting them up front avoids a slow back-and-forth once
    an admin actually starts the real-world registration."""
    company_pan = company_pan.strip().upper()
    if not PAN_PATTERN.fullmatch(company_pan):
        raise HTTPException(status_code=422, detail="Company PAN must be in the format AAAAA9999A.")
    if company_gst:
        company_gst = company_gst.strip().upper()
        # GST characters 3-12 always embed the PAN it was issued against -- a cheap, high-value
        # sanity check that catches a mismatched/mistyped GST without needing a full checksum
        # implementation.
        if len(company_gst) != 15 or company_gst[2:12] != company_pan:
            raise HTTPException(status_code=422, detail="Company GST must be 15 characters and contain the company PAN.")
    authorized_person_aadhar = authorized_person_aadhar.strip().replace(" ", "")
    if not AADHAR_PATTERN.fullmatch(authorized_person_aadhar):
        raise HTTPException(status_code=422, detail="Aadhar number must be exactly 12 digits.")
    contact_number = contact_number.strip()
    if not MOBILE_PATTERN.fullmatch(contact_number):
        raise HTTPException(status_code=422, detail="Contact number must be a valid 10-digit Indian mobile number.")
    contact_email = contact_email.strip()
    if "@" not in contact_email or "." not in contact_email.rsplit("@", 1)[-1]:
        raise HTTPException(status_code=422, detail="Enter a valid email address.")
    if not company_name.strip() or not authorized_signatory_name.strip():
        raise HTTPException(status_code=422, detail="Company name and authorized signatory name are required.")
    try:
        entity = resolve_user_entity(db, user)
        fees = resolve_channel_fees(db, "sms")
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    if not documents:
        raise HTTPException(status_code=422, detail="Upload at least one supporting document")
    combined = float(fees.dlt_platform_fee) + float(fees.dlt_service_fee)
    gst_amount = round(combined * GST_RATE, 2)
    request_row = DltOnboardingRequest(
        entity_id=entity.id,
        status="pending_payment",
        notes=notes,
        company_name=company_name.strip(),
        company_pan=company_pan,
        company_gst=company_gst,
        authorized_signatory_name=authorized_signatory_name.strip(),
        contact_number=contact_number,
        contact_email=contact_email,
        authorized_person_aadhar_encrypted=encrypt_secret(authorized_person_aadhar),
        dlt_platform_fee=fees.dlt_platform_fee,
        dlt_service_fee=fees.dlt_service_fee,
        gst_amount=gst_amount,
        total_amount=round(combined + gst_amount, 2),
    )
    db.add(request_row); db.flush()
    doc_rows = []
    for doc in documents:
        stored_path, _ = _save_upload(doc, f"dlt-requests/{request_row.id}")
        doc_row = DltOnboardingRequestDocument(request_id=request_row.id, filename=doc.filename or "document", stored_path=stored_path, document_type="supporting")
        db.add(doc_row)
        doc_rows.append(doc_row)
    letter_path, _ = _save_upload(authorization_letter, f"dlt-requests/{request_row.id}")
    letter_row = DltOnboardingRequestDocument(request_id=request_row.id, filename=authorization_letter.filename or "authorization_letter", stored_path=letter_path, document_type="authorization_letter")
    db.add(letter_row)
    doc_rows.append(letter_row)
    db.commit(); db.refresh(request_row)
    return _dlt_request_out(request_row, doc_rows)


def _mark_dlt_request_paid(db: Session, request_row: DltOnboardingRequest, reference: str) -> None:
    request_row.status = "paid"
    request_row.payment_reference = reference
    entity = db.get(Entity, request_row.entity_id)
    invoice = create_draft_invoice(
        db, entity, type="dlt_fee",
        base_amount=float(request_row.dlt_platform_fee) + float(request_row.dlt_service_fee),
        gst_amount=float(request_row.gst_amount), reference=request_row.id,
    )
    issue_invoice(db, invoice)
    db.commit()


@router.post("/dlt/request/{request_id}/recharge", response_model=DltOnboardingRequestOut)
def simulate_dlt_request_payment(request_id: str, user: User = Depends(require_user), db: Session = Depends(get_db)):
    try:
        entity = resolve_user_entity(db, user)
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    request_row = db.scalar(select(DltOnboardingRequest).where(DltOnboardingRequest.id == request_id, DltOnboardingRequest.entity_id == entity.id))
    if not request_row:
        raise HTTPException(status_code=404, detail="DLT registration request not found")
    if request_row.status != "pending_payment":
        raise HTTPException(status_code=409, detail=f"This request has already been {request_row.status}")
    _mark_dlt_request_paid(db, request_row, "self-service-dev-payment")
    docs = db.scalars(select(DltOnboardingRequestDocument).where(DltOnboardingRequestDocument.request_id == request_row.id)).all()
    return _dlt_request_out(request_row, docs)


@router.post("/dlt/request/{request_id}/razorpay/order", response_model=RazorpayOrderResponse)
def create_dlt_request_order(request_id: str, user: User = Depends(require_user), db: Session = Depends(get_db)):
    client = _razorpay_client()
    try:
        entity = resolve_user_entity(db, user)
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    request_row = db.scalar(select(DltOnboardingRequest).where(DltOnboardingRequest.id == request_id, DltOnboardingRequest.entity_id == entity.id))
    if not request_row:
        raise HTTPException(status_code=404, detail="DLT registration request not found")
    if request_row.status != "pending_payment":
        raise HTTPException(status_code=409, detail=f"This request has already been {request_row.status}")
    amount_paise = int(round(float(request_row.total_amount) * 100))
    try:
        order = client.order.create({"amount": amount_paise, "currency": "INR", "receipt": f"dlt-request-{request_row.id}", "notes": {"entity_id": entity.id, "request_id": request_row.id}})
    except razorpay.errors.BadRequestError as exc:
        raise HTTPException(status_code=422, detail=f"Razorpay rejected the order request: {exc}") from exc
    db.add(PaymentOrder(entity_id=entity.id, provider="razorpay", provider_order_id=order["id"], amount=request_row.total_amount, purpose="dlt_request", reference_id=request_row.id, status="created"))
    db.commit()
    return RazorpayOrderResponse(order_id=order["id"], key_id=settings.razorpay_key_id, amount_paise=amount_paise)


@router.post("/dlt/request/{request_id}/razorpay/verify", response_model=DltOnboardingRequestOut)
def verify_dlt_request_payment(request_id: str, payload: RazorpayVerifyRequest, user: User = Depends(require_user), db: Session = Depends(get_db)):
    client = _razorpay_client()
    try:
        entity = resolve_user_entity(db, user)
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    request_row = db.scalar(select(DltOnboardingRequest).where(DltOnboardingRequest.id == request_id, DltOnboardingRequest.entity_id == entity.id))
    if not request_row:
        raise HTTPException(status_code=404, detail="DLT registration request not found")
    order = db.scalar(select(PaymentOrder).where(PaymentOrder.provider_order_id == payload.razorpay_order_id, PaymentOrder.entity_id == entity.id, PaymentOrder.reference_id == request_row.id))
    if not order:
        raise HTTPException(status_code=404, detail="No matching payment order for this request")
    if order.status != "created":
        raise HTTPException(status_code=409, detail=f"This order has already been {order.status}")
    try:
        client.utility.verify_payment_signature({
            "razorpay_order_id": payload.razorpay_order_id,
            "razorpay_payment_id": payload.razorpay_payment_id,
            "razorpay_signature": payload.razorpay_signature,
        })
    except razorpay.errors.SignatureVerificationError as exc:
        order.status = "failed"
        db.commit()
        raise HTTPException(status_code=422, detail="Payment signature verification failed") from exc
    order.status = "paid"
    _mark_dlt_request_paid(db, request_row, payload.razorpay_payment_id)
    docs = db.scalars(select(DltOnboardingRequestDocument).where(DltOnboardingRequestDocument.request_id == request_row.id)).all()
    return _dlt_request_out(request_row, docs)


@router.get("/dlt/requests", response_model=list[DltOnboardingRequestOut])
def list_my_dlt_requests(user: User = Depends(require_user), db: Session = Depends(get_db)):
    try:
        entity = resolve_user_entity(db, user)
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    requests = db.scalars(select(DltOnboardingRequest).where(DltOnboardingRequest.entity_id == entity.id).order_by(DltOnboardingRequest.created_at.desc())).all()
    out = []
    for r in requests:
        docs = db.scalars(select(DltOnboardingRequestDocument).where(DltOnboardingRequestDocument.request_id == r.id)).all()
        out.append(_dlt_request_out(r, docs))
    return out


@router.get("/pe-ids")
def list_my_pe_ids(user: User = Depends(require_user), db: Session = Depends(get_db)):
    try:
        entity = resolve_user_entity(db, user)
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    pes = db.scalars(select(PeId).where(PeId.entity_id == entity.id)).all()
    return [{"id": p.id, "value": p.value, "operator": p.operator, "status": p.status} for p in pes]


@router.get("/headers")
def list_my_headers(user: User = Depends(require_user), db: Session = Depends(get_db)):
    try:
        entity = resolve_user_entity(db, user)
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    headers = db.scalars(select(Header).join(PeId).where(PeId.entity_id == entity.id)).all()
    return [{"id": h.id, "pe_id": h.pe_id, "header_id": h.header_id, "value": h.value, "status": h.status} for h in headers]


@router.post("/pe-ids")
def create_my_pe_id(value: str = Form(...), operator: str = Form(...), user: User = Depends(require_capability("channels:manage")), db: Session = Depends(get_db)):
    try:
        entity = resolve_user_entity(db, user)
        require_channel_active(db, entity.id, "sms")
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    if db.scalar(select(PeId).where(PeId.entity_id == entity.id, PeId.value == value)):
        raise HTTPException(status_code=409, detail="This PE ID is already registered for your account")
    pe = PeId(entity_id=entity.id, value=value, operator=operator, status=Status.active)
    db.add(pe)
    try:
        db.commit()
    except IntegrityError:
        db.rollback()
        raise HTTPException(status_code=409, detail="This PE ID is already registered for your account")
    db.refresh(pe)
    return {"id": pe.id, "value": pe.value, "status": pe.status}


@router.post("/pe-ids/{pe_id}/headers")
def create_my_header(pe_id: str, header_id: str = Form(...), value: str = Form(...), user: User = Depends(require_capability("channels:manage")), db: Session = Depends(get_db)):
    try:
        entity = resolve_user_entity(db, user)
        require_channel_active(db, entity.id, "sms")
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    pe = db.get(PeId, pe_id)
    if not pe or pe.entity_id != entity.id:
        raise HTTPException(status_code=404, detail="PE ID not found")
    if db.scalar(select(Header).where(Header.pe_id == pe_id, Header.header_id == header_id)):
        raise HTTPException(status_code=409, detail="This DLT header id is already registered for this PE ID")
    header = Header(pe_id=pe_id, header_id=header_id, value=value, status=Status.active)
    db.add(header)
    try:
        db.commit()
    except IntegrityError:
        db.rollback()
        raise HTTPException(status_code=409, detail="This DLT header id is already registered for this PE ID")
    db.refresh(header)
    return {"id": header.id, "value": header.value, "status": header.status}


@router.post("/templates")
def create_my_template(
    pe_id: str = Form(...),
    header_id: str = Form(...),
    alias: str = Form(...),
    dlt_template_id: str = Form(...),
    body: str = Form(...),
    category: str = Form(default="transactional"),
    user: User = Depends(require_capability("channels:manage")),
    db: Session = Depends(get_db),
):
    try:
        entity = resolve_user_entity(db, user)
        require_channel_active(db, entity.id, "sms")
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    pe, header = db.get(PeId, pe_id), db.get(Header, header_id)
    if not pe or pe.entity_id != entity.id or not header or header.pe_id != pe.id:
        raise HTTPException(status_code=422, detail="Template PE/Header mapping is invalid")
    item = Template(entity_id=entity.id, pe_id=pe_id, header_id=header_id, alias=alias, dlt_template_id=dlt_template_id, body=body, category=category, status=Status.active)
    db.add(item)
    try:
        db.commit()
    except IntegrityError:
        db.rollback()
        raise HTTPException(status_code=409, detail=f"A template with alias '{alias}' or DLT template id '{dlt_template_id}' already exists for this entity")
    db.refresh(item)
    return {"id": item.id, "alias": item.alias, "dlt_template_id": item.dlt_template_id, "category": item.category}


@router.get("/api-keys", response_model=list[ApiKeyOut])
def list_my_api_keys(user: User = Depends(require_recent_2fa), db: Session = Depends(get_db)):
    try:
        entity = resolve_user_entity(db, user)
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    keys = db.scalars(select(ApiKey).where(ApiKey.entity_id == entity.id)).all()
    return [ApiKeyOut(id=k.id, active=k.active, allowed_ips=k.allowed_ips or []) for k in keys]


@router.post("/api-keys/request-otp", response_model=ApiKeyOtpResponse)
def request_api_key_otp(payload: ApiKeyOtpRequest, user: User = Depends(require_recent_2fa), _perm: User = Depends(require_capability("channels:manage")), db: Session = Depends(get_db)):
    """A fresh code gating either sensitive API-key action -- independent of TOTP 2FA (which
    require_recent_2fa above only enforces if the account has it enabled at all): every account
    gets this checkpoint regardless of whether they've set up an authenticator app."""
    return _issue_api_key_otp(db, user, payload.action)


@router.post("/api-keys", response_model=ApiKeyCreateResponse)
def create_my_api_key(payload: ApiKeyCreateRequest, user: User = Depends(require_recent_2fa), _perm: User = Depends(require_capability("channels:manage")), db: Session = Depends(get_db)):
    try:
        entity = resolve_user_entity(db, user)
        require_channel_active(db, entity.id, "sms")
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    _consume_api_key_otp(db, user.id, "generate_key", payload.otp_code)
    raw_key = "pk_live_" + secrets.token_urlsafe(32)
    item = ApiKey(entity_id=entity.id, key_hash=hash_api_key(raw_key), active=True)
    db.add(item); db.commit(); db.refresh(item)
    return ApiKeyCreateResponse(id=item.id, api_key=raw_key)


@router.put("/api-keys/{key_id}/ip-whitelist", response_model=ApiKeyOut)
def update_my_api_key_ip_whitelist(key_id: str, payload: ApiKeyIpWhitelistUpdate, user: User = Depends(require_recent_2fa), _perm: User = Depends(require_capability("channels:manage")), db: Session = Depends(get_db)):
    try:
        entity = resolve_user_entity(db, user)
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    key = db.get(ApiKey, key_id)
    if not key or key.entity_id != entity.id:
        raise HTTPException(status_code=404, detail="API key not found")
    _consume_api_key_otp(db, user.id, "update_ip_whitelist", payload.otp_code)
    key.allowed_ips = [ip.strip() for ip in payload.allowed_ips if ip.strip()]
    db.commit(); db.refresh(key)
    return ApiKeyOut(id=key.id, active=key.active, allowed_ips=key.allowed_ips or [])


@router.post("/api-keys/{key_id}/revoke", response_model=ApiKeyOut)
def revoke_my_api_key(key_id: str, payload: ApiKeyRevokeRequest, user: User = Depends(require_recent_2fa), _perm: User = Depends(require_capability("channels:manage")), db: Session = Depends(get_db)):
    """Soft-revoke (active=False), not a hard delete -- the row (and its key_hash) stays for
    audit purposes, but require_capability-authenticated requests using it will fail the same
    ApiKey.active == True check every other lookup already relies on. There is no un-revoke; a
    revoked key is gone for good, same one-way posture as the plaintext itself."""
    try:
        entity = resolve_user_entity(db, user)
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    key = db.get(ApiKey, key_id)
    if not key or key.entity_id != entity.id:
        raise HTTPException(status_code=404, detail="API key not found")
    _consume_api_key_otp(db, user.id, "revoke_key", payload.otp_code)
    key.active = False
    db.commit(); db.refresh(key)
    return ApiKeyOut(id=key.id, active=key.active, allowed_ips=key.allowed_ips or [])


@router.get("/reports/summary", response_model=ReportsSummaryResponse)
def get_reports_summary(user: User = Depends(require_user), db: Session = Depends(get_db)):
    try:
        entity = resolve_user_entity(db, user)
    except DomainError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    messages = db.scalars(select(Message).where(Message.entity_id == entity.id)).all()
    return ReportsSummaryResponse(
        total=len(messages),
        submitted=sum(1 for m in messages if m.status == "submitted"),
        failed=sum(1 for m in messages if m.status == "failed"),
        accepted=sum(1 for m in messages if m.status == "accepted"),
    )
